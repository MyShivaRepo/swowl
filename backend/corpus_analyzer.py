"""
corpus_analyzer.py — Analyse d'un corpus documentaire → ontologie candidate.

Pour chaque document (chemin local traduit hôte→conteneur, ou URL Web) :
  1. extraction du texte par sections (page PDF, titres HTML/MD, paragraphes DOCX) ;
  2. découpage en chunks porteurs de leur référence {doc, chapter, page} ;
  3. extraction d'éléments OWL + atomes SWRL via l'API Anthropic (Claude).

Le module renvoie une liste de résultats { "ref": {...}, "elements": {...} } ;
la fusion dans l'ontologie connectée + la table de provenance sont faites côté main.py.
"""
from __future__ import annotations

import json
import re
import urllib.request
import urllib.error
from html.parser import HTMLParser
from pathlib import Path

from triple_store import host_to_container

MAX_CHUNK_CHARS = 6000      # taille max d'un chunk envoyé au LLM
MAX_CHUNKS = 40             # garde-fou coût/temps (chunks analysés au total)
ANTHROPIC_URL = "https://api.anthropic.com/v1/messages"
DEFAULT_MODEL = "claude-sonnet-4-6"

# Partie éditable par l'utilisateur (instructions métier, langage naturel)
SYSTEM_PROMPT = (
    "You are an expert ontology-extraction assistant. "
    "From the given document excerpt, extract OWL2 ontology elements following these rules:\n\n"

    "## WHAT TO EXTRACT\n"
    "- **Classes**: concepts explicitly defined or recurring in the text (actors, categories, artefacts, processes, states)\n"
    "- **ObjectProperties**: named relations between two classes (verb linking two nouns that are classes)\n"
    "- **DatatypeProperties**: measurable attributes, dates, identifiers, thresholds with units\n"
    "- **Individuals**: specific named instances of a class (named substances, standards, product types)\n"
    "- **SWRL rules**: ONLY when the text contains an explicit logical constraint or inference "
    "(\"if X then Y\", \"shall not exceed\", numeric thresholds) — never create empty rules\n\n"

    "## NAMING RULES (critical for text highlighting)\n"
    "- IDs always in English CamelCase, singular: `ElectricalDevice` not `ElectricalDevices`\n"
    "- Multi-word concepts → ONE single CamelCase id: `HazardousSubstance` not `Hazardous`+`Substance`\n"
    "- Classes: UpperCamelCase (`HomogeneousMaterial`)\n"
    "- ObjectProperties: lowerCamelCase verb (`containsSubstance`, `isManufacturedBy`)\n"
    "- DatatypeProperties: lowerCamelCase + unit if relevant (`maxConcentration_pct`, `expiryDate`)\n"
    "- Individuals: UpperCamelCase proper noun (`Lead`, `ISO9001`)\n"
    "- `label`: exact natural-language form as it appears in the text, including acronyms "
    "(e.g. id=`ElectricalElectronicEquipment`, label=`EEE`)\n\n"

    "## QUALITY RULES\n"
    "- **Fidelity**: extract ONLY what is explicitly in the text — never add domain knowledge\n"
    "- **ObjectProperties**: always set `domain` and `range` when identifiable from the text\n"
    "- **DatatypeProperties**: always set `domain`; use xsd types: xsd:string, xsd:integer, "
    "xsd:decimal, xsd:date, xsd:dateTime, xsd:boolean, xsd:anyURI\n"
    "- **Individuals**: always set `types` (the class they belong to)\n"
    "- **Hierarchy**: use `subClassOf` whenever the text states \"X is a type of Y\" or \"X is a Y\"\n"
    "- Words under 4 letters or generic words (the, and, for, with, use, has) → do NOT create IDs\n"
    "- 20 precise classes beat 100 vague ones — prefer depth over breadth\n\n"

    "## DO NOT\n"
    "- Invent content absent from the excerpt\n"
    "- Create plural IDs (`Substances` → use `Substance`)\n"
    "- Create an ObjectProperty without domain and range if both classes exist in the text\n"
    "- Create an Individual without types\n"
    "- Create empty SWRL rules (body:[] or head:[])\n"
    "- Split a multi-word concept into separate single-word IDs"
)

# Partie technique fixe, toujours ajoutée par le backend (l'utilisateur ne la voit pas)
_FORMAT_SUFFIX = (
    "\n\nReturn ONLY valid minified JSON with EXACTLY this shape (no markdown, no commentary):\n"
    '{"classes":[{"id":"","label":"","comment":"","subClassOf":[]}],'
    '"object_properties":[{"id":"","label":"","domain":[],"range":[]}],'
    '"datatype_properties":[{"id":"","label":"","domain":[],"range":[]}],'
    '"individuals":[{"id":"","label":"","types":[]}],'
    '"swrl_rules":[{"id":"","label":"","comment":"","body":[],"head":[]}]}\n'
    "ids are CamelCase local names, NO spaces, NO prefixes. domain/range/types/subClassOf "
    "reference class or property ids. SWRL atoms: "
    '{"type":"type_atom","var":"?x","class_id":"Foo"} ; '
    '{"type":"property_atom","subject":"?x","property_id":"hasBar","object":"?y"} ; '
    '{"type":"equality_atom","var":"?x","operator":">","value":"0"}.'
)


# ── Extraction de texte par format ──────────────────────────────────────────

class _HTMLToSections(HTMLParser):
    """Découpe un HTML en sections {chapter, text} d'après les titres h1..h4."""
    def __init__(self):
        super().__init__()
        self.sections = [{"chapter": "", "text": ""}]
        self._in_title = False
        self._title_buf = ""
        self._skip = 0

    def handle_starttag(self, tag, attrs):
        if tag in ("script", "style"):
            self._skip += 1
        if tag in ("h1", "h2", "h3", "h4"):
            self._in_title = True
            self._title_buf = ""

    def handle_endtag(self, tag):
        if tag in ("script", "style") and self._skip:
            self._skip -= 1
        if tag in ("h1", "h2", "h3", "h4"):
            self._in_title = False
            self.sections.append({"chapter": self._title_buf.strip(), "text": ""})

    def handle_data(self, data):
        if self._skip:
            return
        if self._in_title:
            self._title_buf += data
        else:
            self.sections[-1]["text"] += data


def _sections_from_html(html: str):
    p = _HTMLToSections()
    p.feed(html)
    out = []
    for s in p.sections:
        txt = re.sub(r"[ \t]*\n[ \t]*", "\n", s["text"]).strip()
        txt = re.sub(r"\n{3,}", "\n\n", txt)
        if txt:
            out.append({"chapter": s["chapter"], "page": None, "text": txt})
    return out


def _sections_from_text(text: str):
    """TXT / Markdown : découpe sur les titres Markdown (#…) ou numérotés (1.2 …)."""
    lines = text.splitlines()
    sections, cur = [], {"chapter": "", "page": None, "text": ""}
    head_re = re.compile(r"^(#{1,6}\s+.+|(\d+(\.\d+)*)\s+\S.+)$")
    for ln in lines:
        if head_re.match(ln.strip()) and cur["text"].strip():
            sections.append(cur)
            cur = {"chapter": ln.strip().lstrip("# ").strip(), "page": None, "text": ""}
        elif head_re.match(ln.strip()):
            cur["chapter"] = ln.strip().lstrip("# ").strip()
        else:
            cur["text"] += ln + "\n"
    if cur["text"].strip():
        sections.append(cur)
    return [s for s in sections if s["text"].strip()] or [{"chapter": "", "page": None, "text": text}]


# En-têtes type directive/loi (Article 5, Annexe II, Chapter III, 1.2 Titre…)
_PDF_HEADING_RE = re.compile(
    r"^\s*("
    r"(?:Article|Annexe?|Chapter|Chapitre|Section|Titre|Title)\s+[0-9IVXLCivxlc]+[A-Za-z]?"
    r"|[0-9]+(?:\.[0-9]+){0,3}\s+\S.{0,70}"
    r")\s*$",
    re.IGNORECASE | re.MULTILINE,
)


def _detect_chapter(text: str):
    """Repère le 1er en-tête de chapitre/article dans le texte d'une page PDF."""
    m = _PDF_HEADING_RE.search(text)
    if m:
        return re.sub(r"\s+", " ", m.group(1)).strip()[:80]
    return None


def _sections_from_pdf(data: bytes):
    from pypdf import PdfReader
    import io
    reader = PdfReader(io.BytesIO(data))
    out = []
    current_chapter = ""  # reporté de page en page jusqu'au prochain en-tête détecté
    for i, page in enumerate(reader.pages):
        try:
            txt = (page.extract_text() or "").strip()
        except Exception:
            txt = ""
        if not txt:
            continue
        ch = _detect_chapter(txt)
        if ch:
            current_chapter = ch
        out.append({"chapter": current_chapter, "page": i + 1, "text": txt})
    return out


def _sections_from_docx(data: bytes):
    import docx
    import io
    doc = docx.Document(io.BytesIO(data))
    sections, cur = [], {"chapter": "", "page": None, "text": ""}
    for para in doc.paragraphs:
        style = (para.style.name or "").lower() if para.style else ""
        if style.startswith("heading") and para.text.strip():
            if cur["text"].strip():
                sections.append(cur)
            cur = {"chapter": para.text.strip(), "page": None, "text": ""}
        elif para.text.strip():
            cur["text"] += para.text + "\n"
    if cur["text"].strip():
        sections.append(cur)
    return sections or [{"chapter": "", "page": None, "text": "\n".join(p.text for p in doc.paragraphs)}]


def _fetch_bytes(location: str) -> tuple[bytes, str]:
    """Renvoie (contenu, ext) pour une URL ou un chemin local."""
    if re.match(r"^https?://", location, re.I):
        req = urllib.request.Request(location, headers={"User-Agent": "SWOWL/1.1"})
        with urllib.request.urlopen(req, timeout=20) as r:
            status = r.status
            data = r.read()
            ctype = r.headers.get("Content-Type", "")
        if status != 200:
            raise ValueError(f"HTTP {status} — the server did not return the document (try the direct URL to the HTML/PDF content)")
        if not data:
            raise ValueError("Empty response — the server returned no content (the URL may require a browser session or generate content asynchronously)")
        ext = Path(location.split("?")[0]).suffix.lower()
        if not ext:
            ext = ".html" if "html" in ctype else (".pdf" if "pdf" in ctype else ".txt")
        return data, ext
    # chemin local → traduction hôte → conteneur
    p = Path(host_to_container(location))
    if not p.exists() or not p.is_file():
        raise FileNotFoundError(f"File not found: {location}")
    return p.read_bytes(), p.suffix.lower()


def extract_sections(location: str):
    """Renvoie une liste de sections {chapter, page, text} pour un document."""
    data, ext = _fetch_bytes(location)
    if ext == ".pdf":
        return _sections_from_pdf(data)
    if ext == ".docx":
        return _sections_from_docx(data)
    text = data.decode("utf-8", "replace")
    if ext in (".html", ".htm"):
        return _sections_from_html(text)
    return _sections_from_text(text)


def _chunks(sections, doc_name):
    """Aplati les sections en chunks (taille bornée) porteurs de leur référence."""
    out = []
    for s in sections:
        text = s["text"].strip()
        ref_base = {"doc": doc_name, "chapter": s.get("chapter", ""), "page": s.get("page")}
        if len(text) <= MAX_CHUNK_CHARS:
            out.append({"ref": ref_base, "text": text})
        else:
            for i in range(0, len(text), MAX_CHUNK_CHARS):
                out.append({"ref": ref_base, "text": text[i:i + MAX_CHUNK_CHARS]})
    return out


# ── Appels LLM ───────────────────────────────────────────────────────────────

# URLs de base pour les providers cloud compatibles OpenAI
_PROVIDER_BASES: dict[str, str] = {
    "openai": "https://api.openai.com",
    "meta":   "https://api.llama.com",
}

def _call_anthropic(api_key: str, model: str, user_text: str, system: str = "") -> str:
    base = (system.strip() if system.strip() else SYSTEM_PROMPT)
    body = json.dumps({
        "model": model or DEFAULT_MODEL,
        "max_tokens": 4096,
        "system": base + _FORMAT_SUFFIX,
        "messages": [{"role": "user", "content": user_text}],
    }).encode("utf-8")
    req = urllib.request.Request(
        ANTHROPIC_URL, data=body, method="POST",
        headers={
            "x-api-key": api_key,
            "anthropic-version": "2023-06-01",
            "content-type": "application/json",
        },
    )
    with urllib.request.urlopen(req, timeout=90) as r:
        data = json.loads(r.read().decode("utf-8", "replace"))
    parts = data.get("content", [])
    return "".join(p.get("text", "") for p in parts if p.get("type") == "text")


def _call_openai_compat(base_url: str, api_key: str, model: str, user_text: str,
                        system: str = "", timeout: int = 120) -> str:
    """Appel générique compatible OpenAI : OpenAI, Meta/Llama Cloud, Ollama local."""
    sys_msg = (system.strip() if system.strip() else SYSTEM_PROMPT) + _FORMAT_SUFFIX
    url = base_url.rstrip("/") + "/v1/chat/completions"
    body = json.dumps({
        "model": model,
        "max_tokens": 4096,
        "messages": [
            {"role": "system", "content": sys_msg},
            {"role": "user",   "content": user_text},
        ],
    }).encode("utf-8")
    headers: dict[str, str] = {"Content-Type": "application/json"}
    if api_key and api_key.lower() != "ollama":
        headers["Authorization"] = f"Bearer {api_key}"
    req = urllib.request.Request(url, data=body, method="POST", headers=headers)
    with urllib.request.urlopen(req, timeout=timeout) as r:
        data = json.loads(r.read().decode("utf-8", "replace"))
    return data["choices"][0]["message"]["content"]


def _parse_json(text: str) -> dict:
    """Extrait le 1er objet JSON d'une réponse (tolère les fences ```json)."""
    text = text.strip()
    if text.startswith("```"):
        text = re.sub(r"^```[a-zA-Z]*\n?|\n?```$", "", text).strip()
    start = text.find("{")
    if start < 0:
        return {}
    depth, end = 0, -1
    for i in range(start, len(text)):
        if text[i] == "{":
            depth += 1
        elif text[i] == "}":
            depth -= 1
            if depth == 0:
                end = i + 1
                break
    try:
        return json.loads(text[start:end]) if end > 0 else {}
    except Exception:
        return {}


MAX_WORKERS = 5  # appels LLM simultanés


def _analyse_chunk(chunk, api_key, model, system_prompt, provider="anthropic", base_url=""):
    """Traite un chunk → (ref, elements | None, error | None)."""
    if not chunk["text"].strip():
        return chunk["ref"], None, None
    is_local = provider == "meta" and bool(base_url.strip())
    try:
        if provider == "anthropic":
            raw = _call_anthropic(api_key, model, chunk["text"], system_prompt)
        else:
            effective_url = base_url.strip() if base_url.strip() else _PROVIDER_BASES.get(provider, "")
            # Modèle local (Ollama) : génération lente → timeout long
            timeout = 600 if is_local else 120
            raw = _call_openai_compat(effective_url, api_key, model, chunk["text"], system_prompt, timeout=timeout)
        elements = _parse_json(raw)
        print(f"[corpus] chunk {chunk['ref']['page']} OK ({len(elements)} keys) [{provider}]", flush=True)
        return chunk["ref"], elements if elements else None, None
    except urllib.error.HTTPError as e:
        detail = ""
        try:
            err = json.loads(e.read().decode("utf-8", "replace"))
            detail = (err.get("error") or {}).get("message", "")
        except Exception:
            pass
        return chunk["ref"], None, f"LLM HTTP {e.code}{(' — ' + detail) if detail else ''}"
    except Exception as e:  # noqa: BLE001
        return chunk["ref"], None, f"{type(e).__name__}: {e}"


def analyse(documents, api_key: str, model: str = DEFAULT_MODEL, system_prompt: str = "",
            provider: str = "anthropic", base_url: str = "", on_chunk=None):
    """Analyse les documents et renvoie (results, errors).

    results : [{ "ref": {doc, chapter, page}, "elements": {...} }]
    errors  : [{ "doc": name, "error": str }]
    """
    from concurrent.futures import ThreadPoolExecutor, as_completed

    fetch_errors = []
    all_chunks = []
    for d in documents:
        name = (d.get("name") or "").strip() or (d.get("location") or "?")
        location = (d.get("location") or "").strip()
        if not location:
            continue
        try:
            sections = extract_sections(location)
            all_chunks.extend(_chunks(sections, name))
        except Exception as e:  # noqa: BLE001
            fetch_errors.append({"doc": name, "error": f"{type(e).__name__}: {e}"})

    truncated = len(all_chunks) > MAX_CHUNKS
    work = [c for c in all_chunks[:MAX_CHUNKS] if c["text"].strip()]
    # Ollama local sérialise les requêtes (1 GPU) → 1 worker = chunks affichés
    # plus vite, un par un. Cloud/API → parallélisme complet.
    is_local = provider == "meta" and bool(base_url.strip())
    workers = 1 if is_local else MAX_WORKERS
    print(f"[corpus] {len(documents)} doc(s) → {len(work)} chunks to process (parallel={workers})", flush=True)

    results, errors = [], list(fetch_errors)
    abort = False
    with ThreadPoolExecutor(max_workers=workers) as pool:
        futures = {pool.submit(_analyse_chunk, c, api_key, model, system_prompt, provider, base_url): c for c in work}
        for fut in as_completed(futures):
            ref, elements, err = fut.result()
            if on_chunk:
                on_chunk(ref, elements, err, futures[fut].get("text", ""))
            if err:
                errors.append({"doc": ref["doc"], "error": err})
                if "401" in err or "403" in err:
                    abort = True
                    pool.shutdown(wait=False, cancel_futures=True)
                    break
            elif elements:
                results.append({"ref": ref, "elements": elements})

    if truncated:
        errors.append({"doc": "*", "error": f"Corpus truncated to {MAX_CHUNKS} chunks."})
    print(f"[corpus] done: {len(results)} results, {len(errors)} errors", flush=True)
    return results, errors
